import docx
import os

doc_path = r'c:\Users\WebkoWuite\DPS\VB Portefeuille Beheer - Documents\Zakelijke Waarden\Aandelen\EU equity\Research\Nederlandse-pensioenfondsen\The Comprehensive Directory of Dutch Pension Funds and Providers.docx'

def inspect_docx(path):
    if not os.path.exists(path):
        print(f"Error: {path} not found.")
        return

    doc = docx.Document(path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")

    # Sample first 5 paragraphs
    print("\n--- First 5 Paragraphs ---")
    for i, para in enumerate(doc.paragraphs[:5]):
        if para.text.strip():
            print(f"P{i}: {para.text[:100]}...")

    # Sample table structure
    if doc.tables:
        print(f"\n--- Table 0 Content (Top 10 rows) ---")
        table = doc.tables[0]
        for i, row in enumerate(table.rows[:10]):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"R{i}: {cells}")

if __name__ == "__main__":
    inspect_docx(doc_path)
