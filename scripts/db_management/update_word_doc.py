import docx
import sqlite3
import os

doc_path = r'c:\Users\WebkoWuite\DPS\VB Portefeuille Beheer - Documents\Zakelijke Waarden\Aandelen\EU equity\Research\Nederlandse-pensioenfondsen\The Comprehensive Directory of Dutch Pension Funds and Providers.docx'
db_path = r'c:\Users\WebkoWuite\DPS\VB Portefeuille Beheer - Documents\Zakelijke Waarden\Aandelen\EU equity\Research\Nederlandse-pensioenfondsen\data\pension_funds.db'

def update_word_doc():
    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found.")
        return

    # Connect to DB and get data
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute('SELECT name, category, aum_euro_bn, equity_allocation_pct, equity_strategy_notes FROM funds ORDER BY aum_euro_bn DESC, name ASC')
    rows = cursor.fetchall()
    conn.close()

    # Open Doc
    doc = docx.Document(doc_path)
    
    # Add Section at the end
    doc.add_page_break()
    doc.add_heading('Analyzed Financial Metrics (2024)', level=1)
    doc.add_paragraph('This section contains the latest financial metrics and strategy notes extracted during the research phase.')

    # Add Table
    table = doc.add_table(rows=1, cols=5)
    try:
        table.style = 'Table Grid'
    except:
        pass # Default style if Table Grid is missing
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fund Name'
    hdr_cells[1].text = 'Category'
    hdr_cells[2].text = 'AUM (€bn)'
    hdr_cells[3].text = 'Equity %'
    hdr_cells[4].text = 'Strategy Notes'

    for name, cat, aum, eq, notes in rows:
        # Skip noise detected in extraction if any
        if 'Pension Insurers' in name: continue
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(name)
        row_cells[1].text = str(cat) if cat else '-'
        row_cells[2].text = str(aum) if aum is not None else '-'
        row_cells[3].text = str(eq) if eq is not None else '-'
        row_cells[4].text = str(notes) if notes else '-'

    # Save
    doc.save(doc_path)
    print(f"Updated Word document: {doc_path}")

if __name__ == "__main__":
    update_word_doc()
